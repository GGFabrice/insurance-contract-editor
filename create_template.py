from pathlib import Path

from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


BASE_DIR = Path(__file__).resolve().parent
TEMPLATE_DIR = BASE_DIR / "templates"
TEMPLATE_FILE = TEMPLATE_DIR / "avenant_template.docx"


def set_cell_shading(cell, fill):
    """Applique une couleur de fond à une cellule."""
    tc_pr = cell._tc.get_or_add_tcPr()

    shd = tc_pr.find(qn("w:shd"))

    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)

    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text, bold=False):
    """Configure proprement le contenu d'une cellule."""
    cell.text = ""

    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = Pt(9)

    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def ajouter_ligne_information(table, libelle, variable):
    """Ajoute une ligne libellé / valeur."""
    cells = table.add_row().cells

    cells[0].text = libelle
    cells[1].text = variable

    cells[0].paragraphs[0].runs[0].bold = True

    return cells


# ============================================================
# CREATION DU DOCUMENT
# ============================================================

TEMPLATE_DIR.mkdir(parents=True, exist_ok=True)

document = Document()


# ============================================================
# MARGES
# ============================================================

section = document.sections[0]

section.top_margin = Cm(1.5)
section.bottom_margin = Cm(1.5)
section.left_margin = Cm(1.8)
section.right_margin = Cm(1.8)


# ============================================================
# POLICE PAR DEFAUT
# ============================================================

styles = document.styles

styles["Normal"].font.name = "Arial"
styles["Normal"].font.size = Pt(10)


# ============================================================
# EN-TETE
# ============================================================

paragraph = document.add_paragraph()

paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

run = paragraph.add_run("AVENANT AU CONTRAT D'ASSURANCE SANTE")

run.bold = True
run.font.size = Pt(16)


paragraph = document.add_paragraph()

paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

run = paragraph.add_run(
    "AVENANT N° {{ numero_avenant }}"
)

run.bold = True
run.font.size = Pt(13)


paragraph = document.add_paragraph()

paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

run = paragraph.add_run(
    "{{ type_avenant }}"
)

run.bold = True
run.font.size = Pt(11)


document.add_paragraph()


# ============================================================
# INFORMATIONS SOUSCRIPTEUR
# ============================================================

paragraph = document.add_paragraph()

run = paragraph.add_run("1. INFORMATIONS DU SOUSCRIPTEUR")

run.bold = True
run.font.size = Pt(12)


table = document.add_table(
    rows=0,
    cols=2
)

table.alignment = WD_TABLE_ALIGNMENT.CENTER

table.style = "Table Grid"

ajouter_ligne_information(
    table,
    "Code souscripteur",
    "{{ code_souscripteur }}"
)

ajouter_ligne_information(
    table,
    "Raison sociale",
    "{{ raison_sociale }}"
)

ajouter_ligne_information(
    table,
    "Type de souscripteur",
    "{{ type_souscripteur }}"
)

ajouter_ligne_information(
    table,
    "Adresse",
    "{{ adresse_souscripteur }}"
)

ajouter_ligne_information(
    table,
    "Téléphone",
    "{{ telephone_souscripteur }}"
)

ajouter_ligne_information(
    table,
    "Email",
    "{{ email_souscripteur }}"
)


document.add_paragraph()


# ============================================================
# INFORMATIONS CONTRAT
# ============================================================

paragraph = document.add_paragraph()

run = paragraph.add_run("2. INFORMATIONS DU CONTRAT")

run.bold = True
run.font.size = Pt(12)


table = document.add_table(
    rows=0,
    cols=2
)

table.alignment = WD_TABLE_ALIGNMENT.CENTER

table.style = "Table Grid"


ajouter_ligne_information(
    table,
    "Compagnie",
    "{{ compagnie }}"
)

ajouter_ligne_information(
    table,
    "Code compagnie",
    "{{ code_compagnie }}"
)

ajouter_ligne_information(
    table,
    "Intermédiaire",
    "{{ intermediaire }}"
)

ajouter_ligne_information(
    table,
    "Code intermédiaire",
    "{{ code_intermediaire }}"
)

ajouter_ligne_information(
    table,
    "Numéro de compte",
    "{{ numero_compte }}"
)

ajouter_ligne_information(
    table,
    "Numéro de police",
    "{{ numero_police }}"
)

ajouter_ligne_information(
    table,
    "Nature du risque",
    "{{ nature_risque }}"
)

ajouter_ligne_information(
    table,
    "Police",
    "{{ police }}"
)

ajouter_ligne_information(
    table,
    "Date d'effet du contrat",
    "{{ date_effet_contrat }}"
)

ajouter_ligne_information(
    table,
    "Date de fin du contrat",
    "{{ date_fin_contrat }}"
)

ajouter_ligne_information(
    table,
    "Échéance annuelle",
    "{{ echeance_annuelle }}"
)

ajouter_ligne_information(
    table,
    "Fractionnement de prime",
    "{{ fractionnement_prime }}"
)


document.add_paragraph()


# ============================================================
# INFORMATIONS AVENANT
# ============================================================

paragraph = document.add_paragraph()

run = paragraph.add_run("3. INFORMATIONS DE L'AVENANT")

run.bold = True
run.font.size = Pt(12)


table = document.add_table(
    rows=0,
    cols=2
)

table.alignment = WD_TABLE_ALIGNMENT.CENTER

table.style = "Table Grid"


ajouter_ligne_information(
    table,
    "Numéro d'avenant",
    "{{ numero_avenant }}"
)

ajouter_ligne_information(
    table,
    "Type d'avenant",
    "{{ type_avenant }}"
)

ajouter_ligne_information(
    table,
    "Période de début",
    "{{ periode_debut }}"
)

ajouter_ligne_information(
    table,
    "Période de fin",
    "{{ periode_fin }}"
)

ajouter_ligne_information(
    table,
    "Date d'effet",
    "{{ date_effet }}"
)

ajouter_ligne_information(
    table,
    "Date d'édition",
    "{{ date_edition }}"
)

ajouter_ligne_information(
    table,
    "Statut",
    "{{ statut }}"
)


document.add_paragraph()


# ============================================================
# LISTE DES ASSURES
# ============================================================

paragraph = document.add_paragraph()

run = paragraph.add_run("4. LISTE DES ASSURES CONCERNES")

run.bold = True
run.font.size = Pt(12)


table_assures = document.add_table(
    rows=2,
    cols=8
)

table_assures.alignment = WD_TABLE_ALIGNMENT.CENTER

table_assures.style = "Table Grid"


headers = [
    "N°",
    "N° assuré",
    "Nom",
    "Prénom",
    "Date naissance",
    "Lien parenté",
    "Collège",
    "Mouvement",
]


for index, header in enumerate(headers):

    cell = table_assures.rows[0].cells[index]

    set_cell_text(
        cell,
        header,
        bold=True
    )

    set_cell_shading(
        cell,
        "D9EAF7"
    )


# Marqueur utilisé par document_generator.py

table_assures.rows[1].cells[0].text = "LISTE_ASSURES"


document.add_paragraph()


# ============================================================
# DETAIL DES COLLEGES
# ============================================================

paragraph = document.add_paragraph()

run = paragraph.add_run("5. DETAIL DES COLLEGES")

run.bold = True
run.font.size = Pt(12)


table_colleges = document.add_table(
    rows=2,
    cols=7
)

table_colleges.alignment = WD_TABLE_ALIGNMENT.CENTER

table_colleges.style = "Table Grid"


headers_colleges = [
    "Collège",
    "Nombre de personnes",
    "Prime / personne",
    "Prime nette",
    "Accessoire",
    "Taxe",
    "Prime totale",
]


for index, header in enumerate(headers_colleges):

    cell = table_colleges.rows[0].cells[index]

    set_cell_text(
        cell,
        header,
        bold=True
    )

    set_cell_shading(
        cell,
        "D9EAF7"
    )


table_colleges.rows[1].cells[0].text = "DETAIL_COLLEGES"


document.add_paragraph()


# ============================================================
# RECAPITULATIF FINANCIER
# ============================================================

paragraph = document.add_paragraph()

run = paragraph.add_run("6. RECAPITULATIF FINANCIER")

run.bold = True
run.font.size = Pt(12)


table = document.add_table(
    rows=0,
    cols=2
)

table.alignment = WD_TABLE_ALIGNMENT.CENTER

table.style = "Table Grid"


ajouter_ligne_information(
    table,
    "Prime nette",
    "{{ prime_nette }}"
)

ajouter_ligne_information(
    table,
    "Accessoire",
    "{{ accessoire }}"
)

ajouter_ligne_information(
    table,
    "Taxe",
    "{{ taxe }}"
)

ajouter_ligne_information(
    table,
    "Prime totale",
    "{{ prime_totale }}"
)


document.add_paragraph()


# ============================================================
# COMMENTAIRE
# ============================================================

paragraph = document.add_paragraph()

run = paragraph.add_run("7. COMMENTAIRE")

run.bold = True
run.font.size = Pt(12)


paragraph = document.add_paragraph()

paragraph.add_run(
    "{{ commentaire }}"
)


document.add_paragraph()


# ============================================================
# SIGNATURE
# ============================================================

paragraph = document.add_paragraph()

paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

run = paragraph.add_run(
    "Fait pour servir et valoir ce que de droit."
)

run.font.size = Pt(10)


document.add_paragraph()

paragraph = document.add_paragraph()

paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

run = paragraph.add_run(
    "Date : {{ date_edition }}"
)

run.font.size = Pt(10)


document.add_paragraph()

paragraph = document.add_paragraph()

paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

run = paragraph.add_run(
    "Signature et cachet"
)

run.bold = True
run.font.size = Pt(10)


# ============================================================
# SAUVEGARDE
# ============================================================

document.save(
    str(TEMPLATE_FILE)
)


print("Modèle Word créé avec succès.")
print(f"Fichier : {TEMPLATE_FILE}")