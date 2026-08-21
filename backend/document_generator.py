from pathlib import Path
from datetime import date
from decimal import Decimal, ROUND_HALF_UP

from docx import Document


# ============================================================
# CONFIGURATION
# ============================================================

BASE_DIR = Path(__file__).resolve().parent.parent

TEMPLATE_DIR = BASE_DIR / "templates"
DOCUMENT_DIR = BASE_DIR / "documents" / "avenants"

TEMPLATE_FILE = TEMPLATE_DIR / "avenant_template.docx"


# ============================================================
# OUTILS
# ============================================================

def formater_date(valeur):
    """
    Transforme une date en format JJ/MM/AAAA.
    """

    if valeur is None:
        return ""

    if isinstance(valeur, date):
        return valeur.strftime("%d/%m/%Y")

    return str(valeur)


def formater_montant(valeur):
    """
    Formate un montant monétaire.

    Exemple :
        1250000.50 -> 1 250 000,50
    """

    if valeur is None:
        return "0,00"

    try:
        montant = Decimal(str(valeur))

        montant = montant.quantize(
            Decimal("0.01"),
            rounding=ROUND_HALF_UP
        )

        texte = f"{montant:,.2f}"

        return (
            texte
            .replace(",", " ")
            .replace(".", ",")
        )

    except Exception:
        return str(valeur)


def valeur_ou_vide(valeur):
    """
    Retourne une chaîne vide lorsque la valeur est None.
    """

    if valeur is None:
        return ""

    return str(valeur)


def calculer_age(date_naissance, date_reference=None):
    """
    Calcule l'âge exact à une date donnée.
    """

    if not date_naissance:
        return ""

    if date_reference is None:
        date_reference = date.today()

    if not isinstance(date_naissance, date):
        return ""

    age = (
        date_reference.year
        - date_naissance.year
        - (
            (date_reference.month, date_reference.day)
            <
            (date_naissance.month, date_naissance.day)
        )
    )

    return str(age)


def calculer_prorata(date_debut, date_fin):
    """
    Calcule le nombre de jours de couverture
    entre deux dates inclusives.
    """

    if not date_debut or not date_fin:
        return 0

    if date_fin < date_debut:
        return 0

    return (date_fin - date_debut).days + 1


# ============================================================
# GESTION DES LIGNES WORD
# ============================================================

def supprimer_ligne_tableau(row):
    """
    Supprime complètement une ligne d'un tableau Word.
    """

    tr = row._tr
    parent = tr.getparent()

    if parent is not None:
        parent.remove(tr)

# ============================================================
# REMPLACEMENT DES VARIABLES
# ============================================================

def remplacer_variables(document, variables):
    """
    Remplace les variables {{ variable }}

    dans :
        - paragraphes
        - cellules de tableaux

    """

    def remplacer_dans_paragraphe(paragraph):

        texte = paragraph.text

        if not texte:
            return

        nouveau_texte = texte

        for cle, valeur in variables.items():

            placeholder_avec_espace = "{{ " + cle + " }}"
            placeholder_sans_espace = "{{" + cle + "}}"

            nouveau_texte = nouveau_texte.replace(
                placeholder_avec_espace,
                valeur
            )

            nouveau_texte = nouveau_texte.replace(
                placeholder_sans_espace,
                valeur
            )

        if nouveau_texte != texte:
            paragraph.text = nouveau_texte

    # --------------------------------------------------------
    # Paragraphes principaux
    # --------------------------------------------------------

    for paragraph in document.paragraphs:
        remplacer_dans_paragraphe(paragraph)

    # --------------------------------------------------------
    # Tableaux
    # --------------------------------------------------------

    for table in document.tables:

        for row in table.rows:

            for cell in row.cells:

                for paragraph in cell.paragraphs:
                    remplacer_dans_paragraphe(paragraph)


# ============================================================
# RECHERCHE D'UN TABLEAU PAR MARQUEUR
# ============================================================

def trouver_tableau_par_marqueur(document, marqueur):
    """
    Recherche un tableau contenant un marqueur.
    """

    for table in document.tables:

        for row in table.rows:

            for cell in row.cells:

                if marqueur in cell.text:
                    return table

    return None


# ============================================================
# TABLEAU DES ASSURES
# ============================================================

def remplir_tableau_assures(document, avenant):
    """
    Recherche le tableau contenant :

        LISTE_ASSURES

    puis ajoute les assurés concernés par l'avenant.
    """

    tableau_cible = trouver_tableau_par_marqueur(
        document,
        "LISTE_ASSURES"
    )

    if tableau_cible is None:
        return

    # --------------------------------------------------------
    # Recherche et suppression de la ligne marqueur
    # --------------------------------------------------------

    ligne_marqueur = None

    for row in tableau_cible.rows:

        texte_ligne = " ".join(
            cell.text
            for cell in row.cells
        )

        if "LISTE_ASSURES" in texte_ligne:
            ligne_marqueur = row
            break

    if ligne_marqueur is not None:
        supprimer_ligne_tableau(ligne_marqueur)

    # --------------------------------------------------------
    # Récupération des mouvements
    # --------------------------------------------------------

    mouvements = sorted(
        avenant.mouvements,
        key=lambda mouvement: (
            mouvement.mouvement_id or 0
        )
    )

    # --------------------------------------------------------
    # Ajout des assurés
    # --------------------------------------------------------

    compteur = 0

    for mouvement in mouvements:

        assure = mouvement.assure

        if assure is None:
            continue

        compteur += 1

        ligne = tableau_cible.add_row()

        cellules = ligne.cells

        # ----------------------------------------------------
        # N°
        # ----------------------------------------------------

        if len(cellules) > 0:
            cellules[0].text = str(compteur)

        # ----------------------------------------------------
        # Numéro assuré
        # ----------------------------------------------------

        if len(cellules) > 1:
            cellules[1].text = valeur_ou_vide(
                assure.numero_assure
            )

        # ----------------------------------------------------
        # Nom
        # ----------------------------------------------------

        if len(cellules) > 2:
            cellules[2].text = valeur_ou_vide(
                assure.nom
            )

        # ----------------------------------------------------
        # Prénom
        # ----------------------------------------------------

        if len(cellules) > 3:
            cellules[3].text = valeur_ou_vide(
                assure.prenom
            )

        # ----------------------------------------------------
        # Date naissance
        # ----------------------------------------------------

        if len(cellules) > 4:
            cellules[4].text = formater_date(
                assure.date_naissance
            )

        # ----------------------------------------------------
        # Lien parenté
        # ----------------------------------------------------

        if len(cellules) > 5:
            cellules[5].text = valeur_ou_vide(
                assure.lien_parente
            )

        # ----------------------------------------------------
        # Collège
        # ----------------------------------------------------

        if len(cellules) > 6:

            if assure.college:

                cellules[6].text = (
                    f"{assure.college.numero_college} - "
                    f"{assure.college.libelle}"
                )

            else:
                cellules[6].text = ""

        # ----------------------------------------------------
        # Mouvement
        # ----------------------------------------------------

        if len(cellules) > 7:

            cellules[7].text = valeur_ou_vide(
                mouvement.type_mouvement
            )

    # --------------------------------------------------------
    # Si aucun assuré
    # --------------------------------------------------------

    if compteur == 0:

        ligne = tableau_cible.add_row()

        if len(ligne.cells) > 0:

            ligne.cells[0].text = (
                "Aucun assuré concerné par cet avenant."
            )


# ============================================================
# TABLEAU DES COLLEGES
# ============================================================

def remplir_tableau_colleges(document, avenant):
    """
    Recherche le tableau contenant :

        DETAIL_COLLEGES

    puis ajoute les lignes de détail.
    """

    tableau_cible = trouver_tableau_par_marqueur(
        document,
        "DETAIL_COLLEGES"
    )

    if tableau_cible is None:
        return

    # --------------------------------------------------------
    # Recherche ligne marqueur
    # --------------------------------------------------------

    ligne_marqueur = None

    for row in tableau_cible.rows:

        texte_ligne = " ".join(
            cell.text
            for cell in row.cells
        )

        if "DETAIL_COLLEGES" in texte_ligne:
            ligne_marqueur = row
            break

    # --------------------------------------------------------
    # Suppression ligne marqueur
    # --------------------------------------------------------

    if ligne_marqueur is not None:
        supprimer_ligne_tableau(ligne_marqueur)

    # --------------------------------------------------------
    # Récupération des lignes
    # --------------------------------------------------------

    lignes = sorted(
        avenant.lignes,
        key=lambda ligne: (
            ligne.college.numero_college
            if ligne.college
            else 0
        )
    )

    # --------------------------------------------------------
    # Ajout
    # --------------------------------------------------------

    for ligne_avenant in lignes:

        ligne = tableau_cible.add_row()

        cellules = ligne.cells

        college = ligne_avenant.college

        # ----------------------------------------------------
        # Collège
        # ----------------------------------------------------

        if len(cellules) > 0:

            if college:

                cellules[0].text = (
                    f"{college.numero_college} - "
                    f"{college.libelle}"
                )

            else:
                cellules[0].text = ""

        # ----------------------------------------------------
        # Nombre personnes
        # ----------------------------------------------------

        if len(cellules) > 1:

            cellules[1].text = str(
                ligne_avenant.nombre_personnes
            )

        # ----------------------------------------------------
        # Prime / personne
        # ----------------------------------------------------

        if len(cellules) > 2:

            cellules[2].text = formater_montant(
                ligne_avenant.prime_nette_par_personne
            )

        # ----------------------------------------------------
        # Prime nette
        # ----------------------------------------------------

        if len(cellules) > 3:

            cellules[3].text = formater_montant(
                ligne_avenant.prime_nette
            )

        # ----------------------------------------------------
        # Accessoire
        # ----------------------------------------------------

        if len(cellules) > 4:

            cellules[4].text = formater_montant(
                ligne_avenant.accessoire
            )

        # ----------------------------------------------------
        # Taxe
        # ----------------------------------------------------

        if len(cellules) > 5:

            cellules[5].text = formater_montant(
                ligne_avenant.taxe
            )

        # ----------------------------------------------------
        # Total
        # ----------------------------------------------------

        if len(cellules) > 6:

            cellules[6].text = formater_montant(
                ligne_avenant.prime_totale
            )


# ============================================================
# CONSTRUCTION DES VARIABLES
# ============================================================

def construire_variables(avenant):
    """
    Construit le dictionnaire de variables
    utilisé dans le modèle Word.
    """

    contrat = avenant.contrat

    souscripteur = contrat.souscripteur

    type_avenant = valeur_ou_vide(
        avenant.type_avenant
    )

    # --------------------------------------------------------
    # Type avenant
    # --------------------------------------------------------

    if type_avenant == "AVENANT_INCORPORATION":

        libelle_type = "INCORPORATION"

    elif type_avenant == "AVENANT_RETRAIT":

        libelle_type = "RETRAIT"

    else:

        libelle_type = type_avenant

    # --------------------------------------------------------
    # Variables
    # --------------------------------------------------------

    variables = {

        # ====================================================
        # AVENANT
        # ====================================================

        "numero_avenant": valeur_ou_vide(
            avenant.numero_avenant
        ),

        "type_avenant": libelle_type,

        "periode_debut": formater_date(
            avenant.periode_debut
        ),

        "periode_fin": formater_date(
            avenant.periode_fin
        ),

        "date_effet": formater_date(
            avenant.date_effet
        ),

        "date_edition": formater_date(
            avenant.date_edition
        ),

        "statut": valeur_ou_vide(
            avenant.statut
        ),

        "commentaire": valeur_ou_vide(
            avenant.commentaire
        ),

        # ====================================================
        # SOUSCRIPTEUR
        # ====================================================

        "code_souscripteur": valeur_ou_vide(
            souscripteur.code_souscripteur
        ),

        "raison_sociale": valeur_ou_vide(
            souscripteur.raison_sociale
        ),

        "type_souscripteur": valeur_ou_vide(
            souscripteur.type_souscripteur
        ),

        "adresse_souscripteur": valeur_ou_vide(
            souscripteur.adresse
        ),

        "telephone_souscripteur": valeur_ou_vide(
            souscripteur.telephone
        ),

        "email_souscripteur": valeur_ou_vide(
            souscripteur.email
        ),

        # ====================================================
        # CONTRAT
        # ====================================================

        "compagnie": valeur_ou_vide(
            contrat.compagnie
        ),

        "code_compagnie": valeur_ou_vide(
            contrat.code_compagnie
        ),

        "intermediaire": valeur_ou_vide(
            contrat.intermediaire
        ),

        "code_intermediaire": valeur_ou_vide(
            contrat.code_intermediaire
        ),

        "numero_compte": valeur_ou_vide(
            contrat.numero_compte
        ),

        "numero_police": valeur_ou_vide(
            contrat.numero_police
        ),

        "nature_risque": valeur_ou_vide(
            contrat.nature_risque
        ),

        "police": valeur_ou_vide(
            contrat.police
        ),

        "numero_intermediaire_police": valeur_ou_vide(
            contrat.numero_intermediaire_police
        ),

        "date_effet_contrat": formater_date(
            contrat.date_effet
        ),

        "date_fin_contrat": formater_date(
            contrat.date_fin
        ),

        "echeance_annuelle": formater_date(
            contrat.echeance_annuelle
        ),

        "fractionnement_prime": valeur_ou_vide(
            contrat.fractionnement_prime
        ),

        # ====================================================
        # FINANCIER
        # ====================================================

        "prime_nette": formater_montant(
            avenant.prime_nette
        ),

        "accessoire": formater_montant(
            avenant.accessoire
        ),

        "taxe": formater_montant(
            avenant.taxe
        ),

        "prime_totale": formater_montant(
            avenant.prime_totale
        ),

        # ====================================================
        # PRORATA
        # ====================================================

        "nombre_jours_periode": str(
            calculer_prorata(
                avenant.periode_debut,
                avenant.periode_fin
            )
        ),
    }

    return variables


# ============================================================
# GENERATION DU DOCUMENT
# ============================================================

def generer_avenant_document(avenant):
    """
    Génère le document Word correspondant à un avenant.

    Retourne le chemin du fichier généré.
    """

    # --------------------------------------------------------
    # Vérification du modèle
    # --------------------------------------------------------

    if not TEMPLATE_FILE.exists():

        raise FileNotFoundError(
            f"Modèle Word introuvable : {TEMPLATE_FILE}"
        )

    # --------------------------------------------------------
    # Vérification avenant
    # --------------------------------------------------------

    if avenant is None:

        raise ValueError(
            "L'avenant fourni est introuvable."
        )

    if avenant.contrat is None:

        raise ValueError(
            "L'avenant n'est associé à aucun contrat."
        )

    if avenant.contrat.souscripteur is None:

        raise ValueError(
            "Le contrat n'est associé à aucun souscripteur."
        )

    # --------------------------------------------------------
    # Création dossier
    # --------------------------------------------------------

    DOCUMENT_DIR.mkdir(
        parents=True,
        exist_ok=True
    )

    # --------------------------------------------------------
    # Chargement modèle
    # --------------------------------------------------------

    document = Document(
        str(TEMPLATE_FILE)
    )

    # --------------------------------------------------------
    # Variables
    # --------------------------------------------------------

    variables = construire_variables(
        avenant
    )

    remplacer_variables(
        document,
        variables
    )

    # --------------------------------------------------------
    # Tableaux
    # --------------------------------------------------------

    remplir_tableau_assures(
        document,
        avenant
    )

    remplir_tableau_colleges(
        document,
        avenant
    )

    # --------------------------------------------------------
    # Type document
    # --------------------------------------------------------

    if avenant.type_avenant == "AVENANT_INCORPORATION":

        type_document = "incorporation"

    elif avenant.type_avenant == "AVENANT_RETRAIT":

        type_document = "retrait"

    else:

        type_document = (
            valeur_ou_vide(
                avenant.type_avenant
            ).lower()
            or "avenant"
        )

    # --------------------------------------------------------
    # Nom fichier
    # --------------------------------------------------------

    nom_fichier = (
        f"avenant_"
        f"{avenant.numero_avenant}_"
        f"{type_document}.docx"
    )

    chemin_fichier = (
        DOCUMENT_DIR / nom_fichier
    )

    # --------------------------------------------------------
    # Sauvegarde
    # --------------------------------------------------------

    document.save(
        str(chemin_fichier)
    )

    return chemin_fichier